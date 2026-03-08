#!/usr/bin/env python3
"""
将Markdown文档转换为Word docx格式
支持标题、段落、表格、图片、代码块
"""
import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def create_style(doc):
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for i in range(1, 5):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = '黑体'
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if i == 1:
            hs.font.size = Pt(22)
        elif i == 2:
            hs.font.size = Pt(16)
        elif i == 3:
            hs.font.size = Pt(14)
        else:
            hs.font.size = Pt(12)


def parse_md_to_docx(md_path, docx_path, resource_dir=None):
    if resource_dir is None:
        resource_dir = os.path.dirname(md_path)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    create_style(doc)

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]
        raw = line.rstrip('\n')

        if raw.startswith('```') and not in_code_block:
            if in_table and table_rows:
                add_table(doc, table_rows)
                table_rows = []
                in_table = False
            in_code_block = True
            code_lines = []
            i += 1
            continue

        if raw.startswith('```') and in_code_block:
            in_code_block = False
            if code_lines:
                add_code_block(doc, code_lines)
            code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(raw)
            i += 1
            continue

        if raw.startswith('|') and '|' in raw[1:]:
            stripped = raw.strip()
            if re.match(r'^\|[\s\-:|]+\|$', stripped):
                i += 1
                continue
            if not in_table:
                if table_rows:
                    add_table(doc, table_rows)
                    table_rows = []
                in_table = True
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table and table_rows:
                add_table(doc, table_rows)
                table_rows = []
                in_table = False

        if raw.strip() == '' or raw.strip() == '---':
            i += 1
            continue

        heading_match = re.match(r'^(#{1,6})\s+(.*)', raw)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            if level <= 4:
                doc.add_heading(text, level=level)
            else:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(11)
            i += 1
            continue

        img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)', raw.strip())
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            full_path = os.path.join(resource_dir, img_path)
            if os.path.exists(full_path):
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(full_path, width=Inches(5.5))
                    if alt_text:
                        cap = doc.add_paragraph()
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = cap.add_run(alt_text)
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(102, 102, 102)
                except Exception as e:
                    p = doc.add_paragraph(f'[图片: {alt_text}] ({img_path})')
            else:
                p = doc.add_paragraph(f'[图片: {alt_text}] ({img_path})')
            i += 1
            continue

        if raw.startswith('- ') or raw.startswith('* '):
            text = raw[2:].strip()
            text = clean_md_inline(text)
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        num_match = re.match(r'^(\d+)\.\s+(.*)', raw)
        if num_match:
            text = num_match.group(2).strip()
            text = clean_md_inline(text)
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        text = raw.strip()
        if text:
            text = clean_md_inline(text)
            p = doc.add_paragraph()
            add_formatted_text(p, text)

        i += 1

    if in_table and table_rows:
        add_table(doc, table_rows)

    doc.save(docx_path)
    size_kb = os.path.getsize(docx_path) // 1024
    print(f'  生成: {docx_path} ({size_kb}KB)')


def clean_md_inline(text):
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    return text


def add_formatted_text(paragraph, text):
    bold_pattern = re.compile(r'\*\*(.+?)\*\*')
    parts = bold_pattern.split(text)
    is_bold = False
    for part in parts:
        if part:
            run = paragraph.add_run(part)
            run.bold = is_bold
        is_bold = not is_bold


def add_code_block(doc, code_lines):
    text = '\n'.join(code_lines)
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(40, 40, 40)
    pf = p.paragraph_format
    from docx.shared import Pt as PtShared
    pf.space_before = PtShared(4)
    pf.space_after = PtShared(4)
    shading = p._element.get_or_add_pPr()
    bg = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F5F5F5'
    })
    shading.append(bg)


def add_table(doc, rows):
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            if ci < num_cols:
                cell = table.cell(ri, ci)
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(cell_text.strip())
                run.font.size = Pt(9)
                if ri == 0:
                    run.bold = True
                    from docx.oxml import OxmlElement
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:val'), 'clear')
                    shading.set(qn('w:color'), 'auto')
                    shading.set(qn('w:fill'), 'D9E2F3')
                    cell._tc.get_or_add_tcPr().append(shading)


if __name__ == '__main__':
    docs_dir = os.path.dirname(os.path.abspath(__file__))
    out_dir = os.path.join(docs_dir, 'word')
    os.makedirs(out_dir, exist_ok=True)

    conversions = [
        ('需求文档/协议自动识别系统需求规格说明书.md', '协议自动识别系统需求规格说明书.docx', '需求文档'),
        ('设计文档/协议自动识别系统设计文档.md', '协议自动识别系统设计文档.docx', '设计文档'),
        ('测试文档/协议自动识别系统测试方案.md', '协议自动识别系统测试方案.docx', '测试文档'),
        ('测试文档/协议自动识别系统测试用例.md', '协议自动识别系统测试用例.docx', '测试文档'),
        ('部署说明/协议自动识别系统部署说明.md', '协议自动识别系统部署说明.docx', '部署说明'),
    ]

    print('开始生成Word文档...\n')
    for md_rel, docx_name, res_rel in conversions:
        md_path = os.path.join(docs_dir, md_rel)
        docx_path = os.path.join(out_dir, docx_name)
        res_dir = os.path.join(docs_dir, res_rel)
        print(f'转换: {md_rel}')
        parse_md_to_docx(md_path, docx_path, res_dir)

    print('\n全部完成！')
